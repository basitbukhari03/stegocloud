"""
generate_report.py
Generates a professional 3-page StegoCloud project report in DOCX format.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_horizontal_line(doc):
    """Insert a thin horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1E3A5F")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

def heading(doc, text, level=1, color="1A2B4C"):
    """Add a styled section heading."""
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13) if level == 1 else Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, indent=False):
    """Add a body paragraph."""
    p    = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def bullet(doc, text, symbol="•"):
    """Add a bullet point."""
    p   = doc.add_paragraph()
    run = p.add_run(f"{symbol}  {text}")
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent   = Cm(0.8)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    return p

def bold_inline(para, label, value):
    """Add a bold label followed by normal text in same paragraph."""
    r1 = para.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = para.add_run(value)
    r2.font.size = Pt(10.5)

# ─── Document Setup ───────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER + INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

# ── University line ──
univ = doc.add_paragraph()
univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = univ.add_run("Department of Computer Science  |  Hackathon Project 2026")
r.font.size  = Pt(9.5)
r.font.color.rgb = RGBColor(100, 120, 150)
r.italic = True
univ.paragraph_format.space_after = Pt(8)

# ── Dark title banner table ──
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style     = "Table Grid"
cell          = tbl.cell(0, 0)
set_cell_bg(cell, "0D1B2A")

cell_para = cell.paragraphs[0]
cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell_para.paragraph_format.space_before = Pt(18)
cell_para.paragraph_format.space_after  = Pt(4)

r1 = cell_para.add_run("StegoCloud\n")
r1.bold = True
r1.font.size  = Pt(28)
r1.font.color.rgb = RGBColor(0, 212, 255)   # cyan

r2 = cell_para.add_run("Steganography-Based Cloud Data Protection System")
r2.bold = False
r2.font.size  = Pt(12)
r2.font.color.rgb = RGBColor(180, 200, 220)

# Sub-line
sub = cell.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(6)
sub.paragraph_format.space_after  = Pt(14)
rs = sub.add_run("AES-256 Encryption  ·  LSB Steganography  ·  TOTP Multi-Factor Authentication")
rs.font.size  = Pt(10)
rs.font.color.rgb = RGBColor(100, 160, 210)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Project Info table ──
info = doc.add_table(rows=4, cols=2)
info.style     = "Table Grid"
info.alignment = WD_TABLE_ALIGNMENT.CENTER

fields = [
    ("Project Title:",    "StegoCloud – Steganography-Based Cloud Data Protection System"),
    ("Algorithm Used:",   "AES-256-CBC Encryption + LSB (Least Significant Bit) Steganography"),
    ("Technology Stack:", "Python 3.13, Flask, PyCryptodome, Pillow, PyOTP, SQLite, Bootstrap 5"),
    ("Submission Date:",  "2026"),
]

for i, (label, value) in enumerate(fields):
    row = info.rows[i]
    # Label cell
    set_cell_bg(row.cells[0], "EBF2FA")
    lp  = row.cells[0].paragraphs[0]
    lr  = lp.add_run(label)
    lr.bold = True
    lr.font.size = Pt(10)
    row.cells[0].width = Cm(4.5)
    # Value cell
    vp  = row.cells[1].paragraphs[0]
    vr  = vp.add_run(value)
    vr.font.size = Pt(10)

doc.add_paragraph()
add_horizontal_line(doc)

# ── 1. Introduction ──
heading(doc, "1.  Introduction")
body(doc,
    "In the digital age, data security has become a critical concern for individuals, corporations, "
    "and governments alike. Traditional encryption methods protect data contents, but they make it "
    "obvious that sensitive information is being transmitted — a fact that can attract unwanted "
    "attention. Steganography addresses this gap by concealing the very existence of a message "
    "inside an innocent-looking medium such as an image."
)
body(doc,
    "StegoCloud combines two powerful security techniques: AES-256-CBC symmetric encryption and "
    "LSB (Least Significant Bit) image steganography, delivered through a modern Flask-based "
    "cloud storage web application. The result is a system where secret data is first encrypted "
    "and then invisibly embedded inside ordinary PNG images, making it undetectable to the human eye "
    "while remaining recoverable only by authorised users who possess the correct password."
)

# ── 2. Problem Statement ──
heading(doc, "2.  Problem Statement")
body(doc,
    "Conventional cloud storage and messaging platforms provide encryption in transit and at rest, "
    "but the metadata (who is communicating, how often, file sizes) remains visible. Furthermore, "
    "a single security layer can be compromised. The key challenges addressed by StegoCloud are:"
)
for pt in [
    "Sensitive data stored in plain-text or single-layer encrypted files is vulnerable to breaches.",
    "Standard encryption reveals that a secret message exists, inviting targeted attacks.",
    "Weak or absent authentication mechanisms allow unauthorised access to cloud-stored files.",
    "Lack of audit trails makes it impossible to detect and investigate security incidents.",
]:
    bullet(doc, pt)

# ── 3. Objectives ──
heading(doc, "3.  Project Objectives")
for pt in [
    "Implement AES-256-CBC encryption to ensure data confidentiality before embedding.",
    "Apply LSB steganography to hide encrypted data imperceptibly inside PNG images.",
    "Enforce multi-factor authentication (TOTP) to secure user accounts.",
    "Build a role-based access system distinguishing Administrator and User privileges.",
    "Maintain a full audit log recording IP address, timestamp, and action for every event.",
    "Provide a modern, responsive dark-themed web interface for real-world usability.",
]:
    bullet(doc, pt)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — SYSTEM DESIGN & IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()

heading(doc, "4.  System Architecture")
body(doc,
    "StegoCloud follows a layered MVC (Model-View-Controller) architecture built on the "
    "Flask micro-framework. The system is divided into four primary modules:"
)

# Architecture table
arch = doc.add_table(rows=5, cols=2)
arch.style     = "Table Grid"
arch.alignment = WD_TABLE_ALIGNMENT.CENTER

arch_rows = [
    ("Module",             "Responsibility"),
    ("Authentication",     "User registration, login (password + TOTP), session management, brute-force lockout"),
    ("Encryption Engine",  "AES-256-CBC encrypt/decrypt with SHA-256 password-derived keys via PyCryptodome"),
    ("Steganography Engine","LSB hide/extract using Pillow; capacity calculation; PNG I/O"),
    ("Audit & Admin",      "SQLite audit log, rotating file logger, admin panel with role management"),
]

for i, (col1, col2) in enumerate(arch_rows):
    row = arch.rows[i]
    if i == 0:
        set_cell_bg(row.cells[0], "0D1B2A")
        set_cell_bg(row.cells[1], "0D1B2A")
        for c, t in [(row.cells[0], col1), (row.cells[1], col2)]:
            r = c.paragraphs[0].add_run(t)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 212, 255)
            r.font.size = Pt(10)
    else:
        set_cell_bg(row.cells[0], "EBF2FA")
        r0 = row.cells[0].paragraphs[0].add_run(col1)
        r0.bold = True
        r0.font.size = Pt(10)
        r1 = row.cells[1].paragraphs[0].add_run(col2)
        r1.font.size = Pt(10)

doc.add_paragraph()

# ── 5. Core Algorithms ──
heading(doc, "5.  Core Algorithms")

heading(doc, "5.1  AES-256-CBC Encryption", level=2, color="1E5F8A")
body(doc,
    "The Advanced Encryption Standard (AES) with a 256-bit key in Cipher Block Chaining (CBC) mode "
    "is used to encrypt the secret message before embedding. The user's password is hashed with "
    "SHA-256 to derive the 32-byte key. A random 16-byte Initialisation Vector (IV) is prepended "
    "to the ciphertext to ensure that identical messages produce different encrypted outputs."
)
for pt in [
    "Key Derivation:  SHA-256(password) → 32-byte AES key",
    "Mode:            CBC (Cipher Block Chaining) with PKCS7 padding",
    "Library:         PyCryptodome (pycryptodome)",
    "Output format:   Base64(IV + Ciphertext)",
]:
    bullet(doc, pt, symbol="→")

heading(doc, "5.2  LSB Steganography", level=2, color="1E5F8A")
body(doc,
    "Least Significant Bit (LSB) steganography modifies the lowest-order bit of each colour channel "
    "(R, G, B) in every pixel of the cover image. Since a single LSB change alters a pixel value by "
    "at most 1 out of 255, the visual difference is imperceptible to the human eye. A special "
    "32-bit length header is embedded first, followed by the payload bits."
)
for pt in [
    "Each pixel stores 3 bits (1 per R/G/B channel).",
    "A 32-bit header encodes the total payload length for precise extraction.",
    "Capacity formula:  floor((width × height × 3) / 8) bytes.",
    "Library:  Pillow (PIL) for pixel-level read/write.",
]:
    bullet(doc, pt, symbol="→")

heading(doc, "5.3  TOTP Multi-Factor Authentication", level=2, color="1E5F8A")
body(doc,
    "Every account requires a second factor — a Time-based One-Time Password (TOTP) — after the "
    "password step. A unique 32-character Base32 secret is generated per user at registration, "
    "encoded as a QR code, and scanned into Google Authenticator. Codes refresh every 30 seconds "
    "and are verified server-side using the PyOTP library."
)

# ── 6. Security Features ──
heading(doc, "6.  Security Features")

sec_tbl = doc.add_table(rows=6, cols=2)
sec_tbl.style     = "Table Grid"
sec_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

sec_data = [
    ("Brute-Force Protection",  "Account locked for 15 minutes after 5 consecutive failed login attempts."),
    ("Role-Based Access Control","Admin role: full audit log, user management, role promotion/demotion. "
                                  "User role: own files and operations only."),
    ("Audit Logging",           "Every action (LOGIN, HIDE, EXTRACT, FAILED_LOGIN, LOGOUT, ADMIN) is "
                                 "stored with username, IP address, and UTC timestamp."),
    ("Restricted Audit Access", "Only administrators can view system-wide audit logs. Regular users "
                                 "are redirected with an access-denied message."),
    ("CSRF Protection",         "All POST forms include a Flask-WTF CSRF token to prevent cross-site "
                                 "request forgery attacks."),
    ("Security Headers",        "X-Frame-Options, X-XSS-Protection, X-Content-Type-Options, and "
                                 "Referrer-Policy headers set on every response."),
]

for i, (feat, desc) in enumerate(sec_data):
    set_cell_bg(sec_tbl.rows[i].cells[0], "EBF2FA")
    r0 = sec_tbl.rows[i].cells[0].paragraphs[0].add_run(feat)
    r0.bold = True
    r0.font.size = Pt(10)
    r1 = sec_tbl.rows[i].cells[1].paragraphs[0].add_run(desc)
    r1.font.size = Pt(10)
    sec_tbl.rows[i].cells[0].width = Cm(4.8)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — RESULTS, CONCLUSION, REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()

# ── 7. System Workflow ──
heading(doc, "7.  System Workflow")
body(doc, "The end-to-end workflow for hiding and extracting data is described below:")

steps = [
    ("Step 1 — Register & MFA Setup",
     "A new user registers with a username, email, and strong password. The system generates "
     "a unique TOTP secret and presents a QR code for Google Authenticator. MFA must be "
     "confirmed before the account becomes active."),
    ("Step 2 — Login (Two-Factor)",
     "The user enters credentials (Step 1) and is redirected to the TOTP verification page "
     "(Step 2). A valid 6-digit code from the authenticator app grants access to the dashboard."),
    ("Step 3 — Hide Data",
     "The user uploads a PNG cover image, types the secret message, and provides an encryption "
     "password. The system encrypts the message with AES-256-CBC and embeds the ciphertext "
     "into the image's LSBs, producing a stego PNG download."),
    ("Step 4 — Extract Data",
     "The recipient uploads the stego image and enters the correct password. The system reads "
     "LSBs to reconstruct the ciphertext, then AES-256 decrypts it to reveal the plaintext. "
     "A wrong password produces a clear 'Wrong Password!' error with an access-denied panel."),
    ("Step 5 — Admin Oversight",
     "The administrator can view all system audit logs (user, action, IP, timestamp), "
     "activate/deactivate accounts, and promote users to admin or demote admins to user."),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(title + ":  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(14, 90, 150)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

add_horizontal_line(doc)

# ── 8. Results & Testing ──
heading(doc, "8.  Results & Testing")
body(doc,
    "The system was tested end-to-end across all major user flows. The following results "
    "were confirmed during functional testing:"
)

results = [
    "Registration, TOTP setup, and login all completed successfully for multiple accounts.",
    "AES-256-CBC correctly encrypted and decrypted messages of varying lengths (10–2,000 characters).",
    "LSB steganography produced stego images visually identical to the originals; pixel difference = 0–1 per channel.",
    "Wrong-password extraction correctly triggered the 'Wrong Password! — AES-256 decryption blocked' error panel.",
    "Brute-force lockout activated after exactly 5 failed password attempts.",
    "Admin audit log correctly captured LOGIN, HIDE, EXTRACT, FAILED_LOGIN, LOGOUT, and REGISTER events with IP and timestamp.",
    "Role promotion/demotion worked correctly; last-admin demotion was blocked as designed.",
    "Audit log page returned HTTP 403 (redirect) for non-admin users attempting direct URL access.",
]

for r in results:
    bullet(doc, r, symbol="✔")

doc.add_paragraph()

# ── 9. Potential Applications ──
heading(doc, "9.  Potential Applications")

apps = doc.add_table(rows=2, cols=3)
apps.style     = "Table Grid"
apps.alignment = WD_TABLE_ALIGNMENT.CENTER

app_data = [
    ("☁️  Secure Cloud Storage",    "Hide sensitive files inside images stored on public cloud platforms."),
    ("🏥  Medical Records",         "Protect patient data embedded in medical scan images."),
    ("🏛️  Government Intelligence", "Covert communication channels resistant to traffic analysis."),
    ("🔒  Corporate Security",      "Secure internal document sharing with double-layer protection."),
    ("📨  Secure Messaging",        "Invisible message embedding in social-media-shared images."),
    ("🎓  Digital Watermarking",    "Embed ownership metadata inside images to prove copyright."),
]

for idx, (title, desc) in enumerate(app_data):
    row  = idx // 3
    col  = idx  % 3
    cell = apps.cell(row, col)
    p    = cell.paragraphs[0]
    r1   = p.add_run(title + "\n")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor(14, 90, 150)
    r2   = p.add_run(desc)
    r2.font.size = Pt(9)

doc.add_paragraph()
add_horizontal_line(doc)

# ── 10. Conclusion ──
heading(doc, "10. Conclusion")
body(doc,
    "StegoCloud successfully demonstrates the integration of two independent security disciplines — "
    "cryptography and steganography — into a single, production-ready web application. By encrypting "
    "data with AES-256-CBC before hiding it within image pixels using LSB steganography, the system "
    "provides a dual-layer defence that is both mathematically secure and perceptually transparent."
)
body(doc,
    "The addition of TOTP-based two-factor authentication, role-based access control, brute-force "
    "lockout, CSRF protection, and a complete audit logging system brings the project to an "
    "enterprise-grade security standard. StegoCloud proves that steganography is not merely "
    "an academic concept but a practical tool with real-world applications in cloud security, "
    "healthcare, government, and corporate data protection."
)

# ── 11. References ──
heading(doc, "11. References")
refs = [
    "Daemen, J., & Rijmen, V. (2002). The Design of Rijndael: AES — The Advanced Encryption Standard. Springer.",
    "Fridrich, J. (2009). Steganography in Digital Media: Principles, Algorithms, and Applications. Cambridge University Press.",
    "RFC 6238 — TOTP: Time-Based One-Time Password Algorithm. IETF (2011). https://datatracker.ietf.org/doc/html/rfc6238",
    "Python Software Foundation. (2024). Python 3.13 Documentation. https://docs.python.org/3/",
    "Pallets Projects. (2024). Flask Documentation (3.x). https://flask.palletsprojects.com/",
    "Legrandin, L. (2024). PyCryptodome Documentation. https://pycryptodome.readthedocs.io/",
    "Clark, A. (2024). Pillow (PIL Fork) Documentation. https://pillow.readthedocs.io/",
    "Warner, J. (2024). PyOTP — Python One-Time Password Library. https://pyotp.readthedocs.io/",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"[{i}]  {ref}")
    r.font.size = Pt(9.5)

# ── Footer note ──
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run("© 2026 StegoCloud  ·  Cloud Security · Cryptography · Steganography")
fr.font.size = Pt(8.5)
fr.font.color.rgb = RGBColor(130, 150, 170)
fr.italic = True

# ── Save ──
output = r"e:\Antigravity\stegocloud\StegoCloud_Project_Report.docx"
doc.save(output)
print(f"✅ Report saved → {output}")
